import string

import requests
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docxtpl import DocxTemplate
from docx import Document
from docxcompose.composer import Composer
import os


def render_text(sub, composer):
    tpl = DocxTemplate("templates/exports/snipper_text.docx")
    tpl.render(sub)
    tpl.save("temp.docx")
    composer.append(Document("temp.docx"))


def render_selection(sub, composer):
    tpl = DocxTemplate("templates/exports/snipper_selection.docx")
    tpl.render(sub)
    tpl.save("temp.docx")
    composer.append(Document("temp.docx"))


def render_generic_table(position, title, headers, rows, composer, col_aligns=None, valigns=None, merge_cols=None):
    doc_table = Document()

    paragraph = doc_table.add_paragraph()
    title_text = f"{position}. {title}" if position else title
    run = paragraph.add_run(title_text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)

    if not headers:
        return

    table = doc_table.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    if not col_aligns:
        col_aligns = [WD_ALIGN_PARAGRAPH.JUSTIFY] * len(headers)
    if not valigns:
        valigns = [WD_ALIGN_VERTICAL.CENTER] * len(headers)

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(13)
                r.font.bold = True

    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_val in enumerate(row_data):
            cell = row_cells[i]
            cell.text = str(cell_val)

            cell.vertical_alignment = valigns[i] if i < len(valigns) else WD_ALIGN_VERTICAL.CENTER

            align = col_aligns[i] if i < len(col_aligns) else WD_ALIGN_PARAGRAPH.JUSTIFY
            for p in cell.paragraphs:
                p.alignment = align
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(13)

    if merge_cols:
        for col_idx in merge_cols:
            start_row = 1
            while start_row < len(table.rows):
                start_cell = table.cell(start_row, col_idx)
                cell_text = start_cell.text
                end_row = start_row

                while end_row + 1 < len(table.rows) and table.cell(end_row + 1, col_idx).text == cell_text:
                    end_row += 1

                if end_row > start_row:
                    merged_cell = start_cell.merge(table.cell(end_row, col_idx))

                    merged_cell.text = cell_text

                    merged_cell.vertical_alignment = valigns[col_idx] if valigns and col_idx < len(
                        valigns) else WD_ALIGN_VERTICAL.CENTER
                    align = col_aligns[col_idx] if col_aligns and col_idx < len(
                        col_aligns) else WD_ALIGN_PARAGRAPH.JUSTIFY
                    for p in merged_cell.paragraphs:
                        p.alignment = align
                        for r in p.runs:
                            r.font.name = 'Times New Roman'
                            r.font.size = Pt(13)

                start_row = end_row + 1

    doc_table.save("temp_table_generic.docx")
    composer.append(Document("temp_table_generic.docx"))


def render_table(sub, composer):
    render_generic_table(
        position=sub.get('position', ''),
        title=sub.get('sub_title', ''),
        headers=sub.get('headers', []),
        rows=sub.get('rows', []),
        composer=composer
    )


def render_ref_credit(sub, ref_data, composer):
    tpl = DocxTemplate("templates/exports/snipper_credit.docx")
    context = sub.copy()
    context.update({
        "total_credit": ref_data.get("number_theory", 0) + ref_data.get("number_practice", 0),
        "total_theory": ref_data.get("number_theory", 0),
        "total_practice": ref_data.get("number_practice", 0),
        "self_study_hours": ref_data.get("hour_self_study", 0)
    })
    tpl.render(context)
    tpl.save("temp.docx")
    composer.append(Document("temp.docx"))


def render_ref_lecturer_info(sub, ref_data, composer):
    tpl = DocxTemplate("templates/exports/snipper_lecturer_info.docx")
    context = sub.copy()
    context.update({
        "faculty_name": ref_data.get("faculty"),
        "lecturer_name": ref_data.get("last_name") + " " + ref_data.get("first_name"),
        "lecturer_email": ref_data.get("email"),
        "lecturer_room": ref_data.get("room"),
    })
    tpl.render(context)
    tpl.save("temp.docx")
    composer.append(Document("temp.docx"))


def render_ref_subject_requirements(sub, ref_data, composer, jwt_token=None):
    headers = ["STT/No.", "Môn học điều kiện/\nRequirements", "Mã môn học/Code"]
    rows = []
    headers_api = {'Authorization': 'Bearer ' + jwt_token}
    db_types = requests.get("http://127.0.0.1:8000/type-requirements/", headers=headers_api).json()
    categories = {item["name"]: [] for item in db_types}

    for i, item in enumerate(ref_data or []):
        req_type = item.get("requirement_type", {}).get("name", "")
        if req_type in categories:
            categories[req_type].append(item)

    stt = 1
    for req_type, items in categories.items():
        if not items:
            rows.append([
                str(stt),
                f"{req_type}\nKhông có",
                "-"
            ])
            stt += 1
        else:
            for item in items:
                sub_name = item.get("subject_name", "")
                rows.append([
                    str(stt),
                    f"{req_type}\n{sub_name}",
                    str(item.get("subject_code", ""))
                ])
                stt += 1

    h_aligns = [
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        WD_ALIGN_PARAGRAPH.CENTER
    ]

    v_aligns = [
        WD_ALIGN_VERTICAL.CENTER,
        WD_ALIGN_VERTICAL.TOP,
        WD_ALIGN_VERTICAL.CENTER
    ]
    render_generic_table(
        position=sub.get('position', ''),
        title=f"{sub.get('sub_title', '')}:",
        headers=headers,
        rows=rows,
        composer=composer,
        col_aligns=h_aligns,
        valigns=v_aligns
    )


def render_ref_objectives(sub, ref_data, composer):
    headers = ["Mục tiêu môn học/ Course objectives", "Mô tả - Description", "CĐR CTĐT phân bổ cho môn học - PLOs"]
    rows = []

    for i, item in enumerate(ref_data or []):
        co_code = f"CO{i + 1}"
        content = item.get("content", "")

        plo_list = []
        for plo in item.get("programme_learning_outcomes", []):
            plo_list.append(plo.get("name", ""))

        plos_str = "\n".join(filter(None, plo_list))

        rows.append([
            str(co_code),
            str(content.strip()),
            str(plos_str)
        ])

        rows.append([
            str(co_code),
            str(content),
            str(plos_str)
        ])
    h_aligns = [
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        WD_ALIGN_PARAGRAPH.CENTER
    ]

    v_aligns = [
        WD_ALIGN_VERTICAL.CENTER,
        WD_ALIGN_VERTICAL.TOP,
        WD_ALIGN_VERTICAL.CENTER
    ]

    render_generic_table(
        position=sub.get('position', ''),
        title=sub.get('sub_title', ''),
        headers=headers,
        rows=rows,
        composer=composer,
        col_aligns=h_aligns,
        valigns=v_aligns,
        merge_cols=[0]
    )


def render_ref_course_learning_outcomes(sub, ref_data, composer):
    headers1 = [
        "Mục tiêu môn học/\nCourse objectives",
        "CĐR môn học\n(CLO)",
        "Mô tả CĐR -Description"
    ]
    rows1 = []

    plo_set = set()
    for co in (ref_data or []):
        for clo in co.get("clos", []):
            for plo in clo.get("plos", []):
                if "plo_id" in plo:
                    plo_set.add(plo["plo_id"])

    plo_id_list = sorted(list(plo_set))
    headers2 = ["CLOs"] + [f"PLO{pid}" for pid in plo_id_list]
    rows2 = []

    clo_idx = 1
    for co in (ref_data or []):
        co_code = f"CO{co.get('position', '')}"

        for clo in co.get("clos", []):
            clo_code = f"CLO{clo_idx}"
            desc = clo.get("content", "")

            rows1.append([str(co_code), str(clo_code), str(desc)])

            plo_mapping = {p.get("plo_id"): p.get("rating", "") for p in clo.get("plos", [])}
            row_matrix = [str(clo_code)]
            for p_id in plo_id_list:
                rating = plo_mapping.get(p_id)
                row_matrix.append(str(rating) if rating not in [None, ""] else "-")

            rows2.append(row_matrix)
            clo_idx += 1

    title_1 = f"{sub.get('sub_title', '')}\nHọc xong môn học này, người học có khả năng"
    render_generic_table(
        position=sub.get('position', ''),
        title=title_1,
        headers=headers1,
        rows=rows1,
        composer=composer,
        col_aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.JUSTIFY],
        valigns=[WD_ALIGN_VERTICAL.CENTER] * 3,
        merge_cols=[0]
    )

    render_generic_table(
        position="",
        title="Ma trận tích hợp giữa chuẩn đầu ra của môn học và chuẩn đầu ra của chương trình đào tạo",
        headers=headers2,
        rows=rows2,
        composer=composer,
        col_aligns=[WD_ALIGN_PARAGRAPH.CENTER] * len(headers2),
        valigns=[WD_ALIGN_VERTICAL.CENTER] * len(headers2)
    )

    doc_legend = Document()

    legend_table = doc_legend.add_table(rows=3, cols=2)
    legend_data = [
        ["1: Không đáp ứng", "4: Đáp ứng nhiều"],
        ["2: Ít đáp ứng", "5: Đáp ứng rất nhiều"],
        ["3: Đáp ứng trung bình", ""]
    ]

    for r_idx, row_vals in enumerate(legend_data):
        for c_idx, val in enumerate(row_vals):
            cell = legend_table.cell(r_idx, c_idx)
            cell.text = val

            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(13)

    doc_legend.save("temp_legend.docx")
    composer.append(Document("temp_legend.docx"))


def render_ref_learning_material(sub, ref_data, composer, jwt_token=None):
    doc = Document()

    p_main = doc.add_paragraph()
    run_main = p_main.add_run(f"{sub.get('position', '')}. {sub.get('sub_title', '')}")
    run_main.font.name = 'Times New Roman'
    run_main.font.size = Pt(13)

    headers = {'Authorization': 'Bearer ' + jwt_token}
    db_types = requests.get("http://127.0.0.1:8000/type-learning-materials/", headers=headers).json()
    categories = {item["name"]: [] for item in db_types}

    for item in (ref_data or []):
        type_info = item.get("type_material", {})
        type_name = type_info.get("name", "Khác")
        mat_name = item.get("name", "")

        categories[type_name].append(mat_name)

    alphabet = list(string.ascii_lowercase)

    for i, (cat_name, items) in enumerate(categories.items()):

        prefix = alphabet[i] if i < len(alphabet) else str(i)

        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run(f"{prefix}.\t{cat_name}")
        run_sub.font.name = 'Times New Roman'
        run_sub.font.size = Pt(13)
        run_sub.italic = True

        if not items:
            p_content = doc.add_paragraph()
            run_c = p_content.add_run("Không có tài liệu")
            run_c.font.name = 'Times New Roman'
            run_c.font.size = Pt(13)
        else:
            for text in items:
                p_content = doc.add_paragraph()
                run_c = p_content.add_run(f"- {text}")
                run_c.font.name = 'Times New Roman'
                run_c.font.size = Pt(13)

    doc.save("temp_material.docx")
    composer.append(Document("temp_material.docx"))

def render_ref_student_assessment(sub, ref_data, composer):
    headers = [
        "Thành phần đánh giá/\nType of assessment\n(1)",
        "Bài đánh giá\nAssessment methods\n(2)",
        "Thời điểm\nAssessment time\n(3)",
        "CĐR môn học/CLOs\n(4)",
        "Tỷ lệ %\nWeight %\n(5)"
    ]
    rows = []
    total_weight = 0

    for type_item in (ref_data or []):
        type_name = type_item.get("type_assessment", {}).get("name", "")
        total_weight_assessment = 0
        for method in type_item.get("assessment_methods", []):
            method_name = method.get("name", "")
            time = method.get("time", "")
            weight = method.get("weight", 0)
            total_weight_assessment += weight
            clos = method.get("course_learning_outcomes", [])
            clo_list = [f"CLO{clo.get('position')}" for clo in clos if clo.get('position')]
            clo_str = "\n".join(clo_list)

            rows.append([
                str(type_name),
                str(method_name),
                str(time),
                str(clo_str),
                str(weight)+ "%"
            ])
        rows.append([
            "",
            "Tổng cộng/ Total",
            "",
            "",
            str(total_weight_assessment) + "%"
        ])
        total_weight += total_weight_assessment
    rows.append([
        "",
        "Tổng cộng/ Total",
        "",
        "",
        str(total_weight) + "%"
    ])
    h_aligns = [
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER
    ]
    v_aligns = [WD_ALIGN_VERTICAL.CENTER] * 5

    render_generic_table(
        position=sub.get('position', ''),
        title=sub.get('sub_title', ''),
        headers=headers,
        rows=rows,
        composer=composer,
        col_aligns=h_aligns,
        valigns=v_aligns,
        merge_cols=[0]
    )

    doc_extra = Document()

    p_sub = doc_extra.add_paragraph()
    run_sub = p_sub.add_run(
        "a) Hình thức – Nội dung – Thời lượng các bài đánh giá/Assessment format, content and time:")
    run_sub.font.name = 'Times New Roman'
    run_sub.font.size = Pt(13)
    run_sub.italic = True

    p_content = doc_extra.add_paragraph()
    run_content = p_content.add_run("[Nội dung chi tiết hình thức đánh giá sẽ được chèn vào đây]")
    run_content.font.name = 'Times New Roman'
    run_content.font.size = Pt(13)

    doc_extra.save("temp_assessment_extra.docx")
    composer.append(Document("temp_assessment_extra.docx"))

def render_ref_teaching_schedule(sub, ref_data, composer):
    doc = Document()

    p_title = doc.add_paragraph()
    run_title = p_title.add_run(f"{sub.get('position', '')}. {sub.get('sub_title', '')}:")
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(13)
    run_title.bold = True

    groups = {}
    for item in (ref_data or []):
        group_name = item.get("schedule_group", {}).get("name", "Chung")
        if group_name not in groups:
            groups[group_name] = []
        groups[group_name].append(item)

    for group_name, items in groups.items():
        p_group = doc.add_paragraph()
        run_group = p_group.add_run(group_name)
        run_group.font.name = 'Times New Roman'
        run_group.font.size = Pt(13)
        run_group.italic = True

        table = doc.add_table(rows=5, cols=11)
        table.style = 'Table Grid'

        def set_header_cell(r1, c1, r2, c2, text):
            cell = table.cell(r1, c1)
            cell.text = text
            if r1 != r2 or c1 != c2:
                cell.merge(table.cell(r2, c2))

        set_header_cell(0, 0, 3, 0, "Tuần/buổi học\nWeek Section")
        set_header_cell(0, 1, 3, 1, "Nội dung\nContent")
        set_header_cell(0, 2, 3, 2, "CĐR môn học\nCLOs")
        set_header_cell(0, 3, 0, 8, "Hoạt động dạy và học/Teaching and learning")
        set_header_cell(0, 9, 3, 9, "Bài đánh giá\nStudent assessment")
        set_header_cell(0, 10, 3, 10, "Tài liệu chính và tài liệu tham khảo\nTextbooks and materials")

        set_header_cell(1, 3, 2, 4, "Tự học/Self-study")
        set_header_cell(1, 5, 1, 6, "Trực tiếp/FTF")
        set_header_cell(1, 7, 1, 8, "Trực tuyến (nếu có)/Online (if any)")

        set_header_cell(2, 5, 2, 6, "Lý thuyết/Theory")
        set_header_cell(2, 7, 2, 8, "Lý thuyết/Theory")

        set_header_cell(3, 3, 3, 3, "Hoạt động\nActivity")
        set_header_cell(3, 4, 3, 4, "Số giờ\nPeriods")
        set_header_cell(3, 5, 3, 5, "Hoạt động\nActivity")
        set_header_cell(3, 6, 3, 6, "Số giờ\nPeriods")
        set_header_cell(3, 7, 3, 7, "Hoạt động\nActivity")
        set_header_cell(3, 8, 3, 8, "Số giờ\nPeriods")

        nums = ["(1)", "(2)", "(3)", "(4)", "", "(5)", "", "(6)", "", "(7)", "(8)"]
        for c in range(11):
            table.cell(4, c).text = nums[c]

        total_self = 0
        total_ftf = 0
        total_online = 0

        for item in items:
            row_cells = table.add_row().cells

            row_cells[0].text = str(item.get("session_no", ""))
            row_cells[1].text = item.get("content", "").strip()

            clos = item.get("course_learning_outcomes", [])
            clo_strs = [f"CLO{c.get('position', c.get('id', ''))}" for c in clos]
            row_cells[2].text = "\n".join(clo_strs)

            row_cells[3].text = item.get("self_study_activity", "").strip()
            sh = item.get("self_study_hours", 0)
            row_cells[4].text = str(sh)
            total_self += float(sh) if sh else 0

            row_cells[5].text = item.get("offline_activity", "").strip()
            fh = item.get("offline_hours", 0)
            row_cells[6].text = str(fh)
            total_ftf += float(fh) if fh else 0

            row_cells[7].text = item.get("online_activity", "").strip()
            oh = item.get("online_hours", 0)
            row_cells[8].text = str(oh)
            total_online += float(oh) if oh else 0

            assessments = [a.get("name", "") for a in item.get("assessments", [])]
            row_cells[9].text = "\n".join(f"- {a}" for a in assessments if a)

            materials = [m.get("name", "") for m in item.get("learning_materials", [])]
            row_cells[10].text = "\n".join(f"- {m}" for m in materials if m)

        tot_row = table.add_row().cells
        tot_row[0].text = "Tổng cộng/Total"
        tot_row[0].merge(tot_row[2])

        tot_row[3].text = "x"
        tot_row[4].text = f"{total_self:g}"
        tot_row[5].text = "x"
        tot_row[6].text = f"{total_ftf:g}"
        tot_row[7].text = "x"
        tot_row[8].text = f"{total_online:g}"
        tot_row[9].text = "x"
        tot_row[10].text = "x"

        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                if r_idx < 5:
                    align = WD_ALIGN_PARAGRAPH.CENTER
                elif r_idx == len(table.rows) - 1:
                    align = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    if c_idx in [1, 3, 5, 7, 9, 10]:
                        align = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        align = WD_ALIGN_PARAGRAPH.CENTER

                for p in cell.paragraphs:
                    p.alignment = align
                    for r in p.runs:
                        r.font.name = 'Times New Roman'
                        r.font.size = Pt(11)
                        if r_idx < 5:
                            r.font.bold = True

    doc.save("temp_teaching_schedule.docx")
    composer.append(Document("temp_teaching_schedule.docx"))

def render_reference(sub, composer, jwt_token=None):
    ref_code = sub.get("reference_code")
    ref_data = sub.get("reference_data") or {}

    REF_STRATEGIES = {
        "credit": render_ref_credit,
        "director": render_ref_lecturer_info,
        "requirement_subject": lambda s, rd, c: render_ref_subject_requirements(s, rd, c, jwt_token=jwt_token),
        "objectives_and_outcomes": render_ref_objectives,
        "course_learning_outcomes": render_ref_course_learning_outcomes,
        "learning_material": lambda s, rd, c: render_ref_learning_material(s, rd, c, jwt_token=jwt_token),
        "assessment_method": render_ref_student_assessment,
        "teaching_schedule": render_ref_teaching_schedule
    }

    handler = REF_STRATEGIES.get(ref_code)
    if handler:
        handler(sub, ref_data, composer)


my_token = "Glo2MGBTuQfHB50GLaXfzWKfciBKWg"
SUB_STRATEGIES = {
    "text": render_text,
    "selection": render_selection,
    "table": render_table,
    "reference": lambda sub, composer: render_reference(sub, composer, jwt_token=my_token)
}


def fetch_syllabus_data(api_url, jwt_token):
    headers = {'Authorization': 'Bearer ' + jwt_token}
    response = requests.get(api_url, headers=headers)
    raw_json = response.json()

    la_ma = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"]

    formatted_data = []

    for i, main_sec in enumerate(raw_json.get("main_sections", [])):
        section_data = {
            "stt_la_ma": la_ma[i],
            "tieu_de": main_sec["name"],
            "subs": []
        }

        for sub in main_sec.get("sub_sections", []):
            sub_type = sub["type"]
            mapped_sub = {
                "position": sub["position"],
                "sub_title": sub["name"],
                "type": sub_type
            }

            if sub_type == "text":
                mapped_sub["content"] = sub.get("content") or ""

            elif sub_type == "selection":
                group_id = sub.get("attribute_group_id")

                selected_ids = [item["id"] for item in sub.get("selected_values", [])]

                all_options_response = requests.get(f"http://127.0.0.1:8000/attribute-groups/{group_id}/",
                                                    headers=headers).json()
                options_list = all_options_response.get("attribute_values", [])

                mapped_sub["selections"] = []
                for opt in options_list:
                    icon = "☑" if opt["id"] in selected_ids else "☐"
                    mapped_sub["selections"].append({
                        "check_icon": icon,
                        "name_value": opt["name_value"]
                    })

            elif sub_type == "table":
                schema = sub.get("table_schema", {})
                cols = schema.get("columns", [])
                rows = schema.get("rows", [])

                mapped_sub["headers"] = [c["headerName"] for c in cols]
                mapped_sub["rows"] = [
                    [str(r.get(c["field"], "")) for c in cols]
                    for r in rows
                ]

            elif sub_type == "reference":
                mapped_sub["reference_code"] = sub.get("reference_code")
                mapped_sub["reference_data"] = sub.get("reference_data")

            section_data["subs"].append(mapped_sub)

        formatted_data.append(section_data)

    return formatted_data


master_doc = Document("templates/exports/master_template.docx")
composer = Composer(master_doc)
# ☑ (Đã chọn)
#
# ☐ (Chưa chọn)
#
# ☒ (Đánh dấu X)
#
# ✔ (Dấu tick trơn)
api_url = "http://127.0.0.1:8000/syllabuses/51/"
data = fetch_syllabus_data(api_url=api_url, jwt_token=my_token)

for main in data:
    tpl_main = DocxTemplate("templates/exports/snipper_main_title.docx")
    tpl_main.render(main)
    tpl_main.save("temp.docx")
    composer.append(Document("temp.docx"))
    for sub in main["subs"]:
        handler = SUB_STRATEGIES.get(sub["type"])
        if handler:
            handler(sub, composer)

footer_data = {
    "day": "20", "month": "08", "year": "2026",
    "dean_name": "Nguyễn Văn A",
    "compiler_name": "Trần Văn B"
}

tpl_footer = DocxTemplate("templates/exports/snipper_footer.docx")
tpl_footer.render(footer_data)
tpl_footer.save("temp.docx")
composer.append(Document("temp.docx"))

composer.save("templates/exports/output_test.docx")
print("Xuất thành công")
if os.path.exists("temp.docx"):
    os.remove("temp.docx")
