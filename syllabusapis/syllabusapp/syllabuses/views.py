from datetime import datetime
import io

from django.db import transaction
from django.http import HttpResponse
from django.utils import timezone

from django.db.models.functions import Length
from django.shortcuts import render, get_object_or_404
from django_filters.rest_framework import DjangoFilterBackend
from docx import Document
from docxcompose.composer import Composer
from docxtpl import DocxTemplate
from rest_framework import viewsets, status, generics, parsers, permissions, mixins, filters
from rest_framework.decorators import action, permission_classes
from rest_framework.response import Response
from rest_framework.views import APIView

from handlers import get_strategy_map
from syllabuses import perms
from syllabuses.filters import SyllabusFilter, SubjectFilter, LearningMaterialsFilter, UserFilter, LecturerFilter
from syllabuses.models import User, Syllabus, Faculty, Subject, AttributeGroup, TypeRequirement, \
    ProgrammeLearningOutcome, LearningMaterial, TypeLearningMaterial, TypeAssessment, CourseLearningOutcome, Assessment, \
    ScheduleGroup, Major, TrainingProgram, Lecturer, TemplateSyllabus, TemplateSubSection, TemplateMainSection, \
    TemplateTableSubSection, TemplateSelectionSubSection, TemplateTextSubSection
from syllabuses.paginators import UserPaginator, SyllabusPagination, FacultyPagination, SubjectsPagination, \
    LearningMaterialsPagination, MajorPagination, TrainingPagination, ProgrammeLearningOutcomePagination, \
    LecturerPagination, TemplatePagination
from syllabuses.serializer import UserSerializer, UserDetailSerializer, SyllabusSerializer, FacultySerializer, \
    SubjectSerializer, SyllabusDetailSerializer, AttributeGroupListSerializer, TypeRequirementSerializer, \
    ProgrammeLearningOutcomeSerializer, LearningMaterialSerializer, TypeAssessmentSerializer, ScheduleGroupSerializer, \
    MajorSerializer, TrainingProgramSerializer, SyllabusSimpleSerializer, LecturerBasicSerializer, \
    TemplateSyllabusSerializer, TemplateSyllabusBasicSerializer
from syllabuses.strategies import SUB_SECTION_STRATEGIES, DefaultStrategy


class UserView(mixins.ListModelMixin,
               mixins.CreateModelMixin,
               mixins.UpdateModelMixin,
               mixins.DestroyModelMixin,
               viewsets.GenericViewSet):
    queryset = User.objects.filter(is_active=True)
    pagination_class = UserPaginator
    parser_classes = [parsers.MultiPartParser, parsers.JSONParser]
    filter_backends = [DjangoFilterBackend]
    filterset_class  = UserFilter
    http_method_names = ['get', 'post', 'patch', 'delete']

    def get_serializer_class(self):
        if self.action == 'current_user':
            return UserDetailSerializer
        return UserSerializer
    def get_permissions(self):
        if self.action in ['list', 'create', 'partial_update', 'destroy']:
            return [perms.IsAdmin()]
        if self.action == 'current_user':
            return [permissions.IsAuthenticated()]
        return []

    @action(detail=False, methods=['get', 'patch'], url_path='current-user',
            permission_classes=[permissions.IsAuthenticated])
    def current_user(self, request):
        u = request.user
        if request.method.__eq__('PATCH'):
            serializer = self.get_serializer(u, data=request.data, partial=True)
            serializer.is_valid(raise_exception=True)
            serializer.save()
            return Response(serializer.data, status=status.HTTP_200_OK)
        serializer = self.get_serializer(u)
        return Response(serializer.data, status=status.HTTP_200_OK)


class LecturerView(viewsets.ReadOnlyModelViewSet):
    queryset = Lecturer.objects.select_related('user', 'faculty').filter(user__is_active=True)
    serializer_class = LecturerBasicSerializer
    pagination_class = LecturerPagination
    permission_classes = [permissions.IsAuthenticated]
    filter_backends = [DjangoFilterBackend]
    filterset_class = LecturerFilter


class SyllabusView(viewsets.ModelViewSet):
    queryset = Syllabus.objects.all().order_by("-created_date")
    pagination_class = SyllabusPagination
    filter_backends = [DjangoFilterBackend]
    filterset_class = SyllabusFilter
    def get_serializer_class(self):
        if self.action in ['retrieve', 'update', 'partial_update']:
            return SyllabusDetailSerializer
        return SyllabusSerializer

    def get_permissions(self):
        if self.action == 'create':
            return [perms.IsAdmin()]
        return [permissions.IsAuthenticated()]

    def get_queryset(self):
        query = self.queryset

        user = self.request.user
        if user.is_superuser:
            return query
        return query.filter(lecturer__user=user)

    def perform_update(self, serializer):
        user = self.request.user
        time_str = timezone.localtime().strftime('%H:%M %d/%m/%Y')

        edit_msg = f"{user.last_name} {user.first_name} đã chỉnh sửa lúc {time_str}"

        serializer.save(edit_date=edit_msg)

    @action(detail=False, methods=['patch'], url_path='bulk-update-deadlines')
    def bulk_update_deadlines(self, request):
        ids = request.data.get('ids', [])
        start_date = request.data.get('start_date')
        end_date = request.data.get('end_date')
        if not ids:
            return Response({"err_msg": "Thiếu dữ liệu."}, status=status.HTTP_400_BAD_REQUEST)
        if not start_date:
            return Response({"err_msg": "Thiếu ngày bắt đầu."}, status=status.HTTP_400_BAD_REQUEST)
        if not end_date:
            return Response({'err_msg': 'Thiếu ngày kết thúc'}, status=status.HTTP_400_BAD_REQUEST)
        user = self.request.user
        time_str = timezone.localtime().strftime('%H:%M %d/%m/%Y')

        msg = f"{user.last_name} {user.first_name} đã phân công deadline lúc {time_str}"
        Syllabus.objects.filter(id__in=ids).update(start_date_edition=start_date, end_date_edition=end_date,
                                                   edit_date=msg)
        return Response({"msg": f"Đã cập nhật {len(ids)} đề cương."}, status=status.HTTP_200_OK)

    @action(detail=True, methods=['get'], url_path='clos')
    def get_clos(self, request, pk=None):
        syllabus = self.get_object()
        clos = CourseLearningOutcome.objects.filter(course_objective__syllabus=syllabus)

        data = [{
            "id": clo.id,
            "name": f"CLO{clo.course_objective.position}.{clo.position}" if hasattr(clo,
                                                                                    'position') else f"CLO {clo.id}",
            "content": clo.content
        } for clo in clos]

        return Response(data)

    @action(detail=True, methods=['get'], url_path='assessments')
    def get_assessments(self, request, pk=None):
        syllabus = self.get_object()
        assessments = Assessment.objects.filter(syllabus=syllabus).select_related('type_assessment')
        data = [{
            "id": assessment.id,
            "name": assessment.type_assessment.name
        } for assessment in assessments]

        return Response(data)

    @action(detail=True, methods=['get'], url_path='learning-materials')
    def get_learning_materials(self, request, pk=None):
        syllabus = self.get_object()
        learning_materials = LearningMaterial.objects.filter(syllabuses_rel=syllabus)
        data = [{
            "id": learning_material.id,
            "name": learning_material.name
        }for learning_material in learning_materials]
        return Response(data)


class TemplateSyllabusView(viewsets.ModelViewSet):
    queryset = TemplateSyllabus.objects.prefetch_related('main_sections__sub_sections').all()
    permission_classes = [perms.IsSpecialist | perms.IsAdmin]
    pagination_class = TemplatePagination

    def get_serializer_class(self):
        if self.action in ['retrieve', 'update', 'partial_update']:
            return TemplateSyllabusSerializer
        return TemplateSyllabusBasicSerializer

    @action(detail=True, methods=['post'], url_path='clone')
    @transaction.atomic
    def clone_template(self, request, pk=None):
        """Nhân bản template hiện tại thành bản Draft với version mới"""
        old_template = self.get_object()
        new_name = request.data.get("new_name", old_template.name)
        new_version = request.data.get(
            "new_version",
            f"{old_template.version}_copy"
        )

        if TemplateSyllabus.objects.filter(
                name=new_name,
                version=new_version
        ).exists():
            return Response(
                {"err_msg": "Template với tên và phiên bản này đã tồn tại."},
                status=status.HTTP_400_BAD_REQUEST
            )

        # 1. Clone Template
        new_template = TemplateSyllabus.objects.create(
            name=new_name,
            version=new_version,
            is_active=False,
            parent_id=old_template.id
        )

        # 2. Clone Main Sections & Sub Sections
        for old_main in old_template.main_sections.all():
            new_main = TemplateMainSection.objects.create(
                name=old_main.name,
                template=new_template,
                code=old_main.code,
                position=old_main.position
            )
            for old_sub in old_main.sub_sections.all():
                strategy = SUB_SECTION_STRATEGIES.get(old_sub.type, DefaultStrategy())
                strategy.clone(old_sub, new_main)
        serializer = self.get_serializer(new_template)
        return Response(serializer.data, status=status.HTTP_201_CREATED)


class FacultyView(mixins.ListModelMixin,
               mixins.CreateModelMixin,
               mixins.UpdateModelMixin,
               mixins.DestroyModelMixin,
               viewsets.GenericViewSet):
    queryset = Faculty.objects.all()
    serializer_class = FacultySerializer
    pagination_class = FacultyPagination
    http_method_names = ['get', 'post', 'patch', 'delete']

    def get_permissions(self):
        if self.action in ['create', 'partial_update', 'destroy']:
            return [perms.IsAdmin()]
        if self.action == 'list':
            return [permissions.IsAuthenticated()]
        return []


class SubjectView(mixins.ListModelMixin,
               mixins.CreateModelMixin,
               mixins.UpdateModelMixin,
               mixins.DestroyModelMixin,
               viewsets.GenericViewSet):
    queryset = Subject.objects.all()
    serializer_class = SubjectSerializer
    pagination_class = SubjectsPagination
    http_method_names = ['get', 'post', 'patch', 'delete']

    def get_permissions(self):
        if self.action in ['create', 'partial_update', 'destroy']:
            return [perms.IsAdmin()]
        if self.action == 'list':
            return [permissions.IsAuthenticated()]
        return []
    filter_backends = [DjangoFilterBackend]
    filterset_class = SubjectFilter

class MajorView(mixins.ListModelMixin,
               mixins.CreateModelMixin,
               mixins.UpdateModelMixin,
               mixins.DestroyModelMixin,
               viewsets.GenericViewSet):
    queryset = Major.objects.all()
    serializer_class = MajorSerializer
    pagination_class = MajorPagination
    http_method_names = ['get', 'post', 'patch', 'delete']
    permission_classes = [perms.IsAdmin]

    def create(self, request, *args, **kwargs):
        print("===== REQUEST DATA =====")
        print(request.data)

        serializer = self.get_serializer(data=request.data)

        print("===== INITIAL DATA =====")
        print(serializer.initial_data)

        if serializer.is_valid():
            print("===== VALIDATED DATA =====")
            print(serializer.validated_data)

            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)

        print("===== ERRORS =====")
        print(serializer.errors)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def update(self, request, *args, **kwargs):
        instance = self.get_object()

        print("===== REQUEST DATA =====")
        print(request.data)

        serializer = self.get_serializer(
            instance,
            data=request.data,
            partial=kwargs.get("partial", False)
        )

        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)

        print("===== ERRORS =====")
        print(serializer.errors)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class TrainingProgramView(mixins.ListModelMixin,
               mixins.CreateModelMixin,
               mixins.UpdateModelMixin,
               mixins.DestroyModelMixin,
               viewsets.GenericViewSet):
    queryset = TrainingProgram.objects.all()
    serializer_class = TrainingProgramSerializer
    pagination_class = TrainingPagination
    http_method_names = ['get', 'post', 'patch', 'delete']
    permission_classes = [permissions.IsAuthenticated]


    @action(detail=True, methods=['get'], url_path="syllabuses")
    def get_program_syllabuses(self, request, pk=None):
        user = self.request.user
        program = self.get_object()
        syllabuses = Syllabus.objects.filter(trainingprogramsyllabus__training_program=program)
        if not user.is_superuser:
            syllabuses = syllabuses.filter(
                lecturer__user=user
            )
        q = self.request.query_params.get('q')
        if q:
            syllabuses = syllabuses.filter(name__icontains=q)
        page = self.paginate_queryset(syllabuses)
        if page is not None:
            serializer = SyllabusSimpleSerializer(page, many=True)
            return self.get_paginated_response(serializer.data)

        serializer = SyllabusSimpleSerializer(syllabuses, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)


class AttributeGroupView(viewsets.ViewSet, generics.ListAPIView, generics.RetrieveAPIView):
    queryset = AttributeGroup.objects.all()
    serializer_class = AttributeGroupListSerializer

class TypeRequirementView(viewsets.ViewSet, generics.ListAPIView):
    queryset = TypeRequirement.objects.all()
    serializer_class = TypeRequirementSerializer

class ProgrammeLearningOutcomeView(mixins.ListModelMixin,
               mixins.CreateModelMixin,
               mixins.UpdateModelMixin,
               mixins.DestroyModelMixin,
               viewsets.GenericViewSet):
    queryset = ProgrammeLearningOutcome.objects.all()
    serializer_class = ProgrammeLearningOutcomeSerializer
    pagination_class = ProgrammeLearningOutcomePagination
    http_method_names = ['get', 'post', 'patch', 'delete']

    def get_permissions(self):
        if self.action in ['create', 'partial_update', 'destroy']:
            return [perms.IsAdmin()]
        if self.action == 'list':
            return [permissions.IsAuthenticated()]
        return []

    def get_queryset(self):
        return ProgrammeLearningOutcome.objects.annotate(
            name_len=Length('name')
        ).order_by('name_len', 'name')


class LearningMaterialsView(viewsets.ViewSet, generics.ListAPIView):
    queryset = LearningMaterial.objects.all()
    serializer_class = LearningMaterialSerializer
    pagination_class = LearningMaterialsPagination
    filter_backends = [DjangoFilterBackend]
    filterset_class = LearningMaterialsFilter

class TypeLearningMaterialsView(viewsets.ViewSet, generics.ListAPIView):
    queryset = TypeLearningMaterial.objects.all()
    serializer_class = TypeRequirementSerializer

class TypeAssessmentView(viewsets.ViewSet, generics.ListAPIView):
    queryset = TypeAssessment.objects.all()
    serializer_class = TypeAssessmentSerializer

class ScheduleView(viewsets.ViewSet, generics.ListAPIView):
    queryset = ScheduleGroup.objects.all()
    serializer_class = ScheduleGroupSerializer


def get_formatted_syllabus_data_orm(syllabus_id):
    syllabus = get_object_or_404(Syllabus, pk=syllabus_id)

    raw_json = SyllabusDetailSerializer(syllabus).data

    la_ma = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"]
    formatted_data = []

    for i, main_sec in enumerate(raw_json.get("main_sections", [])):
        section_data = {
            "stt_la_ma": la_ma[i],
            "tieu_de": main_sec["name"],
            "subs": []
        }

        for sub in main_sec.get("sub_sections", []):
            sub_type = sub["type"]
            mapped_sub = {
                "position": sub["position"],
                "sub_title": sub["name"],
                "type": sub_type
            }

            if sub_type == "text":
                mapped_sub["content"] = sub.get("content") or ""
            elif sub_type == "selection":
                mapped_sub["attribute_group_id"] = sub.get("attribute_group_id")
                mapped_sub["selected_values"] = sub.get("selected_values", [])
            elif sub_type == "table":
                schema = sub.get("table_schema", {})
                cols = schema.get("columns", [])
                rows = schema.get("rows", [])
                mapped_sub["headers"] = [c["headerName"] for c in cols]
                mapped_sub["rows"] = [[str(r.get(c["field"], "")) for c in cols] for r in rows]
            elif sub_type == "reference":
                mapped_sub["reference_code"] = sub.get("reference_code")
                mapped_sub["reference_data"] = sub.get("reference_data")

            section_data["subs"].append(mapped_sub)

        formatted_data.append(section_data)

    return formatted_data


class ExportSyllabusDocxView(APIView):
    def get(self, request, syllabus_id):
        try:
            data = get_formatted_syllabus_data_orm(syllabus_id)

            strategies = get_strategy_map()

            master_doc = Document("templates/exports/master_template.docx")
            composer = Composer(master_doc)

            for main in data:
                main_stream = io.BytesIO()
                tpl_main = DocxTemplate("templates/exports/snipper_main_title.docx")
                tpl_main.render(main)
                tpl_main.save(main_stream)
                main_stream.seek(0)

                composer.append(Document(main_stream))

                for sub in main["subs"]:
                    handler = strategies.get(sub["type"])
                    if handler:
                        handler(sub, composer)
            now = datetime.now()

            compiler_name = "Chưa cập nhật"
            dean_name = "Chưa cập nhật"
            footer_data = {
                "day": f"{now.day:02d}",
                "month": f"{now.month:02d}",
                "year": str(now.year),
                "dean_name": compiler_name,
                "compiler_name": dean_name
            }
            footer_stream = io.BytesIO()
            tpl_footer = DocxTemplate("templates/exports/snipper_footer.docx")
            tpl_footer.render(footer_data)
            tpl_footer.save(footer_stream)
            footer_stream.seek(0)

            composer.append(Document(footer_stream))

            final_stream = io.BytesIO()
            composer.save(final_stream)
            final_stream.seek(0)

            response = HttpResponse(
                final_stream.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="Syllabus_{syllabus_id}.docx"'
            return response

        except Exception as e:
            return Response({"detail": f"Lỗi xuất file: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
