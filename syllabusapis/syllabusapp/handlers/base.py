import io

from docx import Document
from docxtpl import DocxTemplate
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

from syllabuses.models import AttributeValue


def render_text(sub, composer):
    tpl = DocxTemplate("templates/exports/snipper_text.docx")
    tpl.render(sub)
    stream = io.BytesIO()
    tpl.save(stream)
    stream.seek(0)

    composer.append(Document(stream))

def render_selection(sub, composer):
    group_id = sub.get("attribute_group_id")
    selected_ids = [item["id"] for item in sub.get("selected_values", [])]

    options_list = AttributeValue.objects.filter(attribute_group_id=group_id).values("id", "name_value")

    context = sub.copy()
    context["selections"] = []
    for opt in options_list:
        icon = "☑" if opt["id"] in selected_ids else "☐"
        context["selections"].append({
            "check_icon": icon,
            "name_value": opt["name_value"]
        })

    tpl = DocxTemplate("templates/exports/snipper_selection.docx")
    tpl.render(context)

    stream = io.BytesIO()
    tpl.save(stream)
    stream.seek(0)
    composer.append(Document(stream))

def render_generic_table(position, title, headers, rows, composer, col_aligns=None, valigns=None, merge_cols=None):
    doc_table = Document()

    paragraph = doc_table.add_paragraph()
    title_text = f"{position}. {title}" if position else title
    run = paragraph.add_run(title_text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)

    if not headers:
        return

    table = doc_table.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    if not col_aligns:
        col_aligns = [WD_ALIGN_PARAGRAPH.JUSTIFY] * len(headers)
    if not valigns:
        valigns = [WD_ALIGN_VERTICAL.CENTER] * len(headers)

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(13)
                r.font.bold = True

    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_val in enumerate(row_data):
            cell = row_cells[i]
            cell.text = str(cell_val)

            cell.vertical_alignment = valigns[i] if i < len(valigns) else WD_ALIGN_VERTICAL.CENTER

            align = col_aligns[i] if i < len(col_aligns) else WD_ALIGN_PARAGRAPH.JUSTIFY
            for p in cell.paragraphs:
                p.alignment = align
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(13)

    if merge_cols:
        for col_idx in merge_cols:
            start_row = 1
            while start_row < len(table.rows):
                start_cell = table.cell(start_row, col_idx)
                cell_text = start_cell.text
                end_row = start_row

                while end_row + 1 < len(table.rows) and table.cell(end_row + 1, col_idx).text == cell_text:
                    end_row += 1

                if end_row > start_row:
                    merged_cell = start_cell.merge(table.cell(end_row, col_idx))

                    merged_cell.text = cell_text

                    merged_cell.vertical_alignment = valigns[col_idx] if valigns and col_idx < len(
                        valigns) else WD_ALIGN_VERTICAL.CENTER
                    align = col_aligns[col_idx] if col_aligns and col_idx < len(
                        col_aligns) else WD_ALIGN_PARAGRAPH.JUSTIFY
                    for p in merged_cell.paragraphs:
                        p.alignment = align
                        for r in p.runs:
                            r.font.name = 'Times New Roman'
                            r.font.size = Pt(13)

                start_row = end_row + 1

    stream = io.BytesIO()
    doc_table.save(stream)
    stream.seek(0)

    composer.append(Document(stream))

def render_table(sub, composer):
    render_generic_table(
        position=sub.get('position', ''),
        title=sub.get('sub_title', ''),
        headers=sub.get('headers', []),
        rows=sub.get('rows', []),
        composer=composer
    )