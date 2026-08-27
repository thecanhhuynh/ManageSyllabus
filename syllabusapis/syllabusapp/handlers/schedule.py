import io

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _build_schedule_table(doc, title, items, is_theory=True):
    p_group = doc.add_paragraph()
    run_group = p_group.add_run(f"{title} - {'Lý thuyết' if is_theory else 'Thực hành'}")
    run_group.font.name, run_group.font.size, run_group.italic = 'Times New Roman', Pt(13), True

    cols = 8 if is_theory else 6
    table = doc.add_table(rows=2, cols=cols)
    table.style = 'Table Grid'

    headers = ["Tuần", "Nội dung", "CĐR", "Tự học (Giờ)", "Trực tuyến (Giờ)"]
    if is_theory:
        headers.extend(["Lý thuyết (Giờ)", "Đánh giá", "Tài liệu"])
    else:
        headers = ["Tuần", "Nội dung", "CĐR", "Thực hành (Giờ)", "Đánh giá", "Tài liệu"]

    for i, text in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = text
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name, r.font.size, r.font.bold = 'Times New Roman', Pt(11), True

    total_main = total_self = total_online = 0
    for item in items:
        row = table.add_row().cells
        row[0].text = str(item.get("session_no", ""))
        row[1].text = item.get("content", "").strip()

        clos = item.get("course_learning_outcomes", [])
        row[2].text = "\n".join([f"CLO{c.get('position', c.get('id', ''))}" for c in clos])

        assessments = "\n".join([f"- {a.get('name', '')}" for a in item.get("assessments", [])])
        materials = "\n".join([f"- {m.get('name', '')}" for m in item.get("learning_materials", [])])

        if is_theory:
            sh = item.get("self_study_hours", 0)
            oh = item.get("online_hours", 0)
            th = item.get("offline_hours", 0)

            row[3].text, row[4].text, row[5].text = str(sh), str(oh), str(th)
            row[6].text, row[7].text = assessments, materials

            total_self += float(sh) if sh else 0
            total_online += float(oh) if oh else 0
            total_main += float(th) if th else 0
        else:
            ph = item.get("offline_hours", 0)
            row[3].text = str(ph)
            row[4].text, row[5].text = assessments, materials
            total_main += float(ph) if ph else 0

    tot_row = table.add_row().cells
    tot_row[0].text = "Tổng"
    tot_row[0].merge(tot_row[2])

    if is_theory:
        tot_row[3].text, tot_row[4].text, tot_row[5].text = f"{total_self:g}", f"{total_online:g}", f"{total_main:g}"
    else:
        tot_row[3].text = f"{total_main:g}"

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name, r.font.size = 'Times New Roman', Pt(11)


def render_ref_teaching_schedule(sub, ref_data, composer):
    doc = Document()
    p_title = doc.add_paragraph()
    run_title = p_title.add_run(f"{sub.get('position', '')}. {sub.get('sub_title', '')}:")
    run_title.font.name, run_title.font.size, run_title.bold = 'Times New Roman', Pt(13), True

    groups = {}
    for item in (ref_data or []):
        g_name = item.get("schedule_group", {}).get("name", "Chung")
        groups.setdefault(g_name, []).append(item)

    for group_name, items in groups.items():
        _build_schedule_table(doc, group_name, items, is_theory=True)
        doc.add_paragraph()
        _build_schedule_table(doc, group_name, items, is_theory=False)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    composer.append(Document(stream))