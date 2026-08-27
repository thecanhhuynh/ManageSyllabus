import io
import string
from docx import Document
from docxtpl import DocxTemplate
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

from syllabuses.models import TypeRequirement, TypeLearningMaterial
from .base import render_generic_table


def render_ref_credit(sub, ref_data, composer):
    tpl = DocxTemplate("templates/exports/snipper_credit.docx")
    context = sub.copy()
    context.update({
        "total_credit": ref_data.get("number_theory", 0) + ref_data.get("number_practice", 0),
        "total_theory": ref_data.get("number_theory", 0),
        "total_practice": ref_data.get("number_practice", 0),
        "self_study_hours": ref_data.get("hour_self_study", 0)
    })
    tpl.render(context)

    stream = io.BytesIO()
    tpl.save(stream)
    stream.seek(0)

    composer.append(Document(stream))

def render_ref_subject_requirements(sub, ref_data, composer, jwt_token=None):
    headers = ["STT/No.", "Môn học điều kiện/\nRequirements", "Mã môn học/Code"]
    rows = []
    db_types = TypeRequirement.objects.all().values("name")
    categories = {item["name"]: [] for item in db_types}

    for i, item in enumerate(ref_data or []):
        req_type = item.get("requirement_type", {}).get("name", "")
        if req_type in categories:
            categories[req_type].append(item)

    stt = 1
    for req_type, items in categories.items():
        if not items:
            rows.append([
                str(stt),
                f"{req_type}\nKhông có",
                "-"
            ])
            stt += 1
        else:
            for item in items:
                sub_name = item.get("subject_name", "")
                rows.append([
                    str(stt),
                    f"{req_type}\n{sub_name}",
                    str(item.get("subject_code", ""))
                ])
                stt += 1

    h_aligns = [
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        WD_ALIGN_PARAGRAPH.CENTER
    ]

    v_aligns = [
        WD_ALIGN_VERTICAL.CENTER,
        WD_ALIGN_VERTICAL.TOP,
        WD_ALIGN_VERTICAL.CENTER
    ]
    render_generic_table(
        position=sub.get('position', ''),
        title=f"{sub.get('sub_title', '')}:",
        headers=headers,
        rows=rows,
        composer=composer,
        col_aligns=h_aligns,
        valigns=v_aligns
    )

def render_ref_objectives(sub, ref_data, composer):
    headers = ["Mục tiêu môn học/ Course objectives", "Mô tả - Description", "CĐR CTĐT phân bổ cho môn học - PLOs"]
    rows = []

    for i, item in enumerate(ref_data or []):
        co_code = f"CO{i + 1}"
        content = item.get("content", "")

        plo_list = []
        for plo in item.get("programme_learning_outcomes", []):
            plo_list.append(plo.get("name", ""))

        plos_str = "\n".join(filter(None, plo_list))

        rows.append([
            str(co_code),
            str(content.strip()),
            str(plos_str)
        ])

        rows.append([
            str(co_code),
            str(content),
            str(plos_str)
        ])
    h_aligns = [
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        WD_ALIGN_PARAGRAPH.CENTER
    ]

    v_aligns = [
        WD_ALIGN_VERTICAL.CENTER,
        WD_ALIGN_VERTICAL.TOP,
        WD_ALIGN_VERTICAL.CENTER
    ]

    render_generic_table(
        position=sub.get('position', ''),
        title=sub.get('sub_title', ''),
        headers=headers,
        rows=rows,
        composer=composer,
        col_aligns=h_aligns,
        valigns=v_aligns,
        merge_cols=[0]
    )

def render_ref_course_learning_outcomes(sub, ref_data, composer):
    headers1 = [
        "Mục tiêu môn học/\nCourse objectives",
        "CĐR môn học\n(CLO)",
        "Mô tả CĐR -Description"
    ]
    rows1 = []

    plo_set = set()
    for co in (ref_data or []):
        for clo in co.get("clos", []):
            for plo in clo.get("plos", []):
                if "plo_id" in plo:
                    plo_set.add(plo["plo_id"])

    plo_id_list = sorted(list(plo_set))
    headers2 = ["CLOs"] + [f"PLO{pid}" for pid in plo_id_list]
    rows2 = []

    clo_idx = 1
    for co in (ref_data or []):
        co_code = f"CO{co.get('position', '')}"

        for clo in co.get("clos", []):
            clo_code = f"CLO{clo_idx}"
            desc = clo.get("content", "")

            rows1.append([str(co_code), str(clo_code), str(desc)])

            plo_mapping = {p.get("plo_id"): p.get("rating", "") for p in clo.get("plos", [])}
            row_matrix = [str(clo_code)]
            for p_id in plo_id_list:
                rating = plo_mapping.get(p_id)
                row_matrix.append(str(rating) if rating not in [None, ""] else "-")

            rows2.append(row_matrix)
            clo_idx += 1

    title_1 = f"{sub.get('sub_title', '')}\nHọc xong môn học này, người học có khả năng"
    render_generic_table(
        position=sub.get('position', ''),
        title=title_1,
        headers=headers1,
        rows=rows1,
        composer=composer,
        col_aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.JUSTIFY],
        valigns=[WD_ALIGN_VERTICAL.CENTER] * 3,
        merge_cols=[0]
    )

    render_generic_table(
        position="",
        title="Ma trận tích hợp giữa chuẩn đầu ra của môn học và chuẩn đầu ra của chương trình đào tạo",
        headers=headers2,
        rows=rows2,
        composer=composer,
        col_aligns=[WD_ALIGN_PARAGRAPH.CENTER] * len(headers2),
        valigns=[WD_ALIGN_VERTICAL.CENTER] * len(headers2)
    )

    doc_legend = Document()

    legend_table = doc_legend.add_table(rows=3, cols=2)
    legend_data = [
        ["1: Không đáp ứng", "4: Đáp ứng nhiều"],
        ["2: Ít đáp ứng", "5: Đáp ứng rất nhiều"],
        ["3: Đáp ứng trung bình", ""]
    ]

    for r_idx, row_vals in enumerate(legend_data):
        for c_idx, val in enumerate(row_vals):
            cell = legend_table.cell(r_idx, c_idx)
            cell.text = val

            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(13)

    stream = io.BytesIO()
    doc_legend.save(stream)
    stream.seek(0)

    composer.append(Document(stream))

def render_ref_learning_material(sub, ref_data, composer, jwt_token=None):
    doc = Document()

    p_main = doc.add_paragraph()
    run_main = p_main.add_run(f"{sub.get('position', '')}. {sub.get('sub_title', '')}")
    run_main.font.name = 'Times New Roman'
    run_main.font.size = Pt(13)

    db_types = TypeLearningMaterial.objects.all().values("name")
    categories = {item["name"]: [] for item in db_types}

    for item in (ref_data or []):
        type_info = item.get("type_material", {})
        type_name = type_info.get("name", "Khác")
        mat_name = item.get("name", "")

        categories[type_name].append(mat_name)

    alphabet = list(string.ascii_lowercase)

    for i, (cat_name, items) in enumerate(categories.items()):

        prefix = alphabet[i] if i < len(alphabet) else str(i)

        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run(f"{prefix}.\t{cat_name}")
        run_sub.font.name = 'Times New Roman'
        run_sub.font.size = Pt(13)
        run_sub.italic = True

        if not items:
            p_content = doc.add_paragraph()
            run_c = p_content.add_run("Không có tài liệu")
            run_c.font.name = 'Times New Roman'
            run_c.font.size = Pt(13)
        else:
            for text in items:
                p_content = doc.add_paragraph()
                run_c = p_content.add_run(f"- {text}")
                run_c.font.name = 'Times New Roman'
                run_c.font.size = Pt(13)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    composer.append(Document(stream))
