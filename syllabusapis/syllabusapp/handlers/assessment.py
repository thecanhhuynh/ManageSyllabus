import io

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from handlers.base import render_generic_table


def render_ref_student_assessment(sub, ref_data, composer):
    headers = [
        "Thành phần đánh giá/\nType of assessment\n(1)",
        "Bài đánh giá\nAssessment methods\n(2)",
        "Thời điểm\nAssessment time\n(3)",
        "CĐR môn học/CLOs\n(4)",
        "Tỷ lệ %\nWeight %\n(5)"
    ]
    rows = []
    total_weight = 0

    for type_item in (ref_data or []):
        type_name = type_item.get("type_assessment", {}).get("name", "")
        total_weight_assessment = 0
        for method in type_item.get("assessment_methods", []):
            method_name = method.get("name", "")
            time = method.get("time", "")
            weight = method.get("weight", 0)
            total_weight_assessment += weight
            clos = method.get("course_learning_outcomes", [])
            clo_list = [f"CLO{clo.get('position')}" for clo in clos if clo.get('position')]
            clo_str = "\n".join(clo_list)

            rows.append([
                str(type_name),
                str(method_name),
                str(time),
                str(clo_str),
                str(weight)+ "%"
            ])
        rows.append([
            "",
            "Tổng cộng/ Total",
            "",
            "",
            str(total_weight_assessment) + "%"
        ])
        total_weight += total_weight_assessment
    rows.append([
        "",
        "Tổng cộng/ Total",
        "",
        "",
        str(total_weight) + "%"
    ])
    h_aligns = [
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER
    ]
    v_aligns = [WD_ALIGN_VERTICAL.CENTER] * 5

    render_generic_table(
        position=sub.get('position', ''),
        title=sub.get('sub_title', ''),
        headers=headers,
        rows=rows,
        composer=composer,
        col_aligns=h_aligns,
        valigns=v_aligns,
        merge_cols=[0]
    )

    doc_extra = Document()

    p_sub = doc_extra.add_paragraph()
    run_sub = p_sub.add_run(
        "a) Hình thức – Nội dung – Thời lượng các bài đánh giá/Assessment format, content and time:")
    run_sub.font.name = 'Times New Roman'
    run_sub.font.size = Pt(13)
    run_sub.italic = True

    p_content = doc_extra.add_paragraph()
    run_content = p_content.add_run("[Nội dung chi tiết hình thức đánh giá sẽ được chèn vào đây]")
    run_content.font.name = 'Times New Roman'
    run_content.font.size = Pt(13)

    stream = io.BytesIO()
    doc_extra.save(stream)
    stream.seek(0)

    composer.append(Document(stream))